import openai
from docx import Document
def remove_newlines(text):
    print("removed")
    text2 = text.replace("/n", "")
    text3 = text2.replace("니다 ", "니다. ")
    return text3.replace("요 ", "요. ")

def split_text_into_paragraphs(text, max_length=800):
    paragraphs = []
    start = 0
    end = 0

    while start < len(text):
        # 800자 내에서 가장 마지막 온점 찾기
        end = start + max_length
        if end >= len(text):
            end = len(text)
        else:
            while end > start and text[end] not in '.!?':
                end -= 1
            # 온점, 물음표, 느낌표 중 하나를 찾았다면 그 뒷문자까지 포함
            end += 1
        if end - start < 100:
            end = start + 700

        paragraph = text[start:end].strip()
        paragraphs.append(paragraph)
        start = end

    return paragraphs


def correct_text(prompt, model="gpt-3.5-turbo"):
    instruction = f"""이상한 부분을 고친 버전을 줘"""

    messages = [
        {"role": "system",
         "content": instruction},  # instruction here
        {"role": "user", "content": prompt}
    ]
    response = openai.ChatCompletion.create(
        model=model,
        messages=messages,
        temperature=0,
    )
    return response.choices[0].message["content"]

def englished_text(prompt, subject='생리학', model="gpt-4"):
    instruction = f"""의학 용어들은 영어로 쓰고 다른 명사나 동사 등은 한글로 써서 고쳐줘"""
    messages = [
        {"role": "system",
         "content": instruction},  # instruction here
        {"role": "user", "content": prompt}
    ]
    response = openai.ChatCompletion.create(
        model=model,
        messages=messages,
        temperature=0,
    )
    return response.choices[0].message["content"]



def topic_splited_text(prompt, subject='생리학', model="gpt-4"):
    instruction = f"""주어진 문단을 적절히 2문단으로 쪼개줘"""
    messages = [
        {"role": "system",
         "content": instruction},  # instruction here
        {"role": "user", "content": prompt}
    ]
    response = openai.ChatCompletion.create(
        model=model,
        messages=messages,
        temperature=0,
    )
    return response.choices[0].message["content"]

def summarised_text(prompt, subject = '생리학',model="gpt-4"):
    instruction = f"""너는 강연을 듣고, 그걸 아주 요약해서 핵심만 전달하는 전문가야. 너는 핵심단어, 매커니즘, 여러 요소들간의 상호작용을 발견하는 능력이 아주 뛰어나. 
내가 보내준 강의 스크립트 받아서 내가 보내주는 문단을 마크다운 형식으로 개조식 스타일로 보여줘. 한국어로 해주고 마지막에 고난이도 객관식 퀴즈 3개와 답을 추가해줘"""
    # instruction에서 한국어 키워드를 제거하는게 좋은듯
    messages = [
        {"role": "system",
         "content": instruction},  # instruction here
        {"role": "user", "content": prompt}
    ]
    response = openai.ChatCompletion.create(
        model=model,
        messages=messages,
        temperature=0,
    )
    return response.choices[0].message["content"]

def concatenate_three_elements(lst):
    result = []
    temp = ""
    for i, elem in enumerate(lst):
        temp += elem + '\n'
        if (i + 1) % 3 == 0:
            result.append(temp[:-1])
            temp = ""
    if temp:
        result.append(temp[:-1])
    return result

def convert_markdown_to_word(filename, outputfilename):
    # Create a new Word Document
    doc = Document()

    # 파일에서 마크다운 컨텐츠 읽기
    with open(filename, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    # 각 라인을 순회하면서 워드 문서에 추가
    for line in lines:
        line = line.strip()  # 공백 제거
        if line.startswith('# '):
            doc.add_heading(line[2:], level=0)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=1)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=2)
        elif line.startswith('- '):
            doc.add_paragraph(line[2:], style='ListBullet')
        else:
            doc.add_paragraph(line)

    # 워드 문서 저장
    doc.save(f'{outputfilename}.docx')
